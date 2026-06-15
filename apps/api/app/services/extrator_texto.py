"""
Extração de texto de PDF e Word, com OCR local via Tesseract para PDFs escaneados.
Dependências: pymupdf (fitz), python-docx, pytesseract, pdf2image, Pillow.
Tesseract deve estar instalado no sistema (apt: tesseract-ocr tesseract-ocr-por).
"""
import logging
import fitz  # PyMuPDF
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

try:
    import pytesseract
    from PIL import Image
    import io
    _TESSERACT_DISPONIVEL = True
except ImportError:
    _TESSERACT_DISPONIVEL = False


def extrair_texto(arquivo_path: str, arquivo_tipo: str) -> tuple[str, bool, int]:
    """
    Returns: (texto, usou_ocr, n_paginas)
    """
    if arquivo_tipo == "pdf":
        return _extrair_pdf(arquivo_path)
    elif arquivo_tipo in ("docx", "doc"):
        return _extrair_docx(arquivo_path), False, 0
    raise ValueError(f"Tipo não suportado: {arquivo_tipo}")


def _extrair_pdf(path: str) -> tuple[str, bool, int]:
    doc = fitz.open(path)
    n_paginas = len(doc)
    partes = []
    usou_ocr = False

    for page in doc:
        txt = page.get_text().strip()
        if len(txt) < 50:
            txt = _ocr_pagina(page)
            if txt:
                usou_ocr = True
        partes.append(txt)

    doc.close()
    return "\n\n".join(partes), usou_ocr, n_paginas


def _ocr_pagina(page) -> str:
    """OCR local via Tesseract. Requer tesseract-ocr instalado no sistema."""
    if not _TESSERACT_DISPONIVEL:
        logger.warning("pytesseract não instalado — página escaneada ignorada")
        return ""
    try:
        pix = page.get_pixmap(dpi=200)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        return pytesseract.image_to_string(img, lang="por+eng")
    except Exception as exc:
        logger.warning("OCR falhou na página: %s", exc)
        return ""


def _extrair_docx(path: str) -> str:
    doc = DocxDocument(path)
    partes = [p.text for p in doc.paragraphs if p.text.strip()]
    # Também extrai tabelas
    for tabela in doc.tables:
        for linha in tabela.rows:
            celulas = [c.text.strip() for c in linha.cells if c.text.strip()]
            if celulas:
                partes.append(" | ".join(celulas))
    return "\n".join(partes)


def dividir_em_chunks(texto: str, max_chars: int = 12000) -> list[str]:
    """
    Divide o texto em chunks respeitando quebras de seção quando possível.
    Editais longos são divididos; contratos curtos ficam em um único chunk.
    """
    if len(texto) <= max_chars:
        return [texto]

    # Divide por parágrafos, agrupando até max_chars
    paragrafos = texto.split("\n\n")
    chunks, atual = [], ""
    for p in paragrafos:
        if len(atual) + len(p) > max_chars and atual:
            chunks.append(atual)
            atual = p
        else:
            atual += "\n\n" + p
    if atual:
        chunks.append(atual)
    return chunks
